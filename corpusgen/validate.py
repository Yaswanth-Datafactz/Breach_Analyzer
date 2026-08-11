"""Corpus validator (docs/plan.md §8): re-opens EVERY generated file with a
real parser (PyMuPDF / python-docx / openpyxl / stdlib email / plain read)
and asserts each planting's value is present at its recorded location;
then asserts scenario quotas, cross-identity value uniqueness (collisions
only where scripted), manifest-count consistency, and prints a
per-file_class table. Non-zero exit on any failure — the generation CLI
runs this automatically, and `python -m corpusgen --validate` runs it
standalone.

Image-borne text (pdf_scanned pages, png screenshots) is verified
OCR-tolerantly: the planted value must appear in the Tesseract text of its
recorded page/image at fuzz.partial_ratio >= 85 — exact match can fail BY
DESIGN on the degraded scans, so the exact-recovery rate is measured and
reported (renderer-tuned to ~80-95%) rather than enforced per planting.
Problem files are verified by their problem kind: password-protected PDFs
must refuse to open bare and open with the recorded password,
zero-byte files must be empty, wrong-extension files must sniff to their
recorded true type (and their recoverable plantings must verify there).
"""

from __future__ import annotations

import hashlib
import io
import json
from collections import Counter, defaultdict
from email import policy
from email.parser import BytesParser
from pathlib import Path

import pikepdf
import pymupdf
import pytesseract
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from rapidfuzz import fuzz

from corpusgen.config import PROFILES

OCR_FUZZ_THRESHOLD = 85  # matches renderers/scanned_pdf.py's render-time bar
OCR_DPI = 200

QUARANTINE_REASON_CODES = {
    "password_protected", "corrupt", "zero_byte", "wrong_extension",
    "unsupported", "ocr_garbage", "parser_error",
}

# Families where a value shared by two identities would corrupt ER ground
# truth. Medical conditions legitimately repeat and are excluded; dob
# repeats are possible and harmless (never a sole identifier).
UNIQUE_FAMILIES = [
    "ssn", "credit_card", "phone", "email", "address", "drivers_license",
    "passport", "financial_account", "username", "password", "employee_id",
]

TRAP_KINDS = [
    "trap_order_number",
    "trap_card_invalid",
    "trap_test_ssn",
    "trap_placeholder",
    "trap_signature_email",
]


def _collapse(text: str) -> str:
    return " ".join(text.split())


def _ocr(img: Image.Image, table: bool = False) -> str:
    config = "--psm 6" if table else ""
    return _collapse(pytesseract.image_to_string(img, config=config))


class _OcrStats:
    """Exact-recovery bookkeeping for the image-borne classes (reported,
    not enforced — the fuzz>=85 bar is the pass/fail criterion)."""

    def __init__(self) -> None:
        self.exact: Counter = Counter()
        self.total: Counter = Counter()

    def record(self, bucket: str, value: str, ocr_text: str) -> None:
        self.total[bucket] += 1
        if _collapse(value) in ocr_text:
            self.exact[bucket] += 1


def _fuzz_check(uid: str, p: dict, ocr_text: str, where: str, errors: list[str]) -> None:
    score = fuzz.partial_ratio(_collapse(p["value"]), ocr_text)
    if score < OCR_FUZZ_THRESHOLD:
        errors.append(
            f"{uid}: {p['element_type']} {p['value']!r} not OCR-recoverable "
            f"{where} (fuzz {score:.0f} < {OCR_FUZZ_THRESHOLD})"
        )


# --- per-class planting checkers (shared by top-level docs, email
# attachments, and unlocked/mislabeled problem files) -----------------------


def _check_pdf_text(pdf: pymupdf.Document, plantings: list[dict], uid: str,
                    errors: list[str]) -> None:
    page_texts = [_collapse(page.get_text()) for page in pdf]
    for p in plantings:
        page = p["location"]["page"]
        if page > len(page_texts) or _collapse(p["value"]) not in page_texts[page - 1]:
            errors.append(f"{uid}: {p['element_type']} {p['value']!r} not on page {page}")


def _check_pdf_scanned(pdf: pymupdf.Document, plantings: list[dict], uid: str,
                       errors: list[str], stats: _OcrStats) -> None:
    page_texts: dict[int, str] = {}
    for p in plantings:
        page = p["location"]["page"]
        if page > len(pdf):
            errors.append(f"{uid}: {p['element_type']} {p['value']!r} page {page} missing")
            continue
        if page not in page_texts:
            pix = pdf[page - 1].get_pixmap(dpi=OCR_DPI)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            page_texts[page] = _ocr(img)
        stats.record("pdf_scanned", p["value"], page_texts[page])
        _fuzz_check(uid, p, page_texts[page], f"on page {page}", errors)


def _check_docx(document: Document, plantings: list[dict], uid: str,
                errors: list[str]) -> None:
    for p in plantings:
        loc = p["location"]
        if "paragraph" in loc:
            idx = loc["paragraph"]
            ok = idx <= len(document.paragraphs) and p["value"] in document.paragraphs[idx - 1].text
        else:
            try:
                cell = document.tables[loc["table"] - 1].rows[loc["row"] - 1].cells[loc["col"] - 1]
                ok = p["value"] in cell.text
            except IndexError:
                ok = False
        if not ok:
            errors.append(f"{uid}: {p['element_type']} {p['value']!r} not at {loc}")


def _check_xlsx(workbook, plantings: list[dict], uid: str, errors: list[str]) -> None:
    for p in plantings:
        loc = p["location"]
        value = workbook[loc["sheet"]][loc["cell"]].value
        if value is None or p["value"] not in str(value):
            errors.append(f"{uid}: {p['element_type']} {p['value']!r} not at {loc}")


def _check_lines(lines: list[str], plantings: list[dict], uid: str,
                 errors: list[str]) -> None:
    for p in plantings:
        line_no = p["location"]["line"]
        if line_no > len(lines) or p["value"] not in lines[line_no - 1]:
            errors.append(f"{uid}: {p['element_type']} {p['value']!r} not on line {line_no}")


def _check_png_image(data: bytes, plantings: list[dict], uid: str,
                     errors: list[str], stats: _OcrStats) -> None:
    ocr_text = _ocr(Image.open(io.BytesIO(data)), table=True)
    for p in plantings:
        stats.record("png", p["value"], ocr_text)
        _fuzz_check(uid, p, ocr_text, f"at {p['location']}", errors)


# --- eml -------------------------------------------------------------------


def _check_attachment_plantings(filename: str, data: bytes, plantings: list[dict],
                                uid: str, errors: list[str]) -> None:
    suffix = Path(filename).suffix
    if suffix == ".docx":
        _check_docx(Document(io.BytesIO(data)), plantings, uid, errors)
    elif suffix == ".xlsx":
        wb = load_workbook(io.BytesIO(data), read_only=True)
        try:
            _check_xlsx(wb, plantings, uid, errors)
        finally:
            wb.close()
    elif suffix == ".pdf":
        with pymupdf.open(stream=data, filetype="pdf") as pdf:
            _check_pdf_text(pdf, plantings, uid, errors)
    else:
        errors.append(f"{uid}: unverifiable attachment type {filename}")


def _check_eml(doc: dict, path: Path, errors: list[str]) -> None:
    uid = doc["doc_uid"]
    with path.open("rb") as fh:
        msg = BytesParser(policy=policy.default).parse(fh)

    body_part = msg.get_body(preferencelist=("plain",))
    body_lines = body_part.get_content().splitlines() if body_part else []
    actual = {
        att.get_filename(): att.get_payload(decode=True)
        for att in msg.iter_attachments()
    }

    recorded = doc["attachments"] or []
    if {a["filename"] for a in recorded} != set(actual):
        errors.append(
            f"{uid}: attachment mismatch — manifest {sorted(a['filename'] for a in recorded)}, "
            f"file {sorted(actual)}"
        )
    for a in recorded:
        data = actual.get(a["filename"])
        if data is None:
            continue
        digest = hashlib.sha256(data).hexdigest()
        if digest != a["sha256"] or len(data) != a["byte_size"]:
            errors.append(f"{uid}: attachment {a['filename']} sha256/size drift")

    by_attachment: dict[str, list[dict]] = defaultdict(list)
    for p in doc["plantings"]:
        part = p["location"]["part"]
        if part == "headers":
            header_value = str(msg[p["location"]["header"]] or "")
            if p["value"] not in header_value:
                errors.append(f"{uid}: {p['value']!r} not in {p['location']['header']} header")
        elif part == "body":
            line_no = p["location"]["line"]
            if line_no > len(body_lines) or p["value"] not in body_lines[line_no - 1]:
                errors.append(
                    f"{uid}: {p['element_type']} {p['value']!r} not on body line {line_no}"
                )
        else:
            by_attachment[part.removeprefix("attachment:")].append(p)

    for filename, plantings in by_attachment.items():
        data = actual.get(filename)
        if data is None:
            errors.append(f"{uid}: plantings reference missing attachment {filename}")
            continue
        _check_attachment_plantings(filename, data, plantings, uid, errors)


# --- problem files ---------------------------------------------------------


def _check_problem(doc: dict, path: Path, errors: list[str], stats: _OcrStats) -> None:
    uid = doc["doc_uid"]
    problem = doc["problem"]
    kind = problem.get("kind")
    data = path.read_bytes()

    for key in ("kind", "expected_reason_code", "recoverable", "recovery_hint"):
        if key not in problem:
            errors.append(f"{uid}: problem entry missing {key!r}")
    if problem.get("expected_reason_code") not in QUARANTINE_REASON_CODES:
        errors.append(f"{uid}: bad reason_code {problem.get('expected_reason_code')!r}")

    if kind == "password_pdf":
        try:
            with pikepdf.open(path):
                errors.append(f"{uid}: password_pdf opened WITHOUT a password")
        except pikepdf.PasswordError:
            pass
        with pymupdf.open(path) as pdf:
            if not pdf.authenticate(problem["password"]):
                errors.append(f"{uid}: recorded password does not unlock the pdf")
            else:
                _check_pdf_text(pdf, doc["plantings"], uid, errors)
    elif kind == "truncated_pdf":
        if not data.startswith(b"%PDF") or b"%%EOF" in data or not data:
            errors.append(f"{uid}: truncated_pdf is not a headless pdf wreck")
        if doc["plantings"]:
            errors.append(f"{uid}: truncated_pdf must record no plantings")
    elif kind == "zero_byte":
        if len(data) != 0:
            errors.append(f"{uid}: zero_byte file has {len(data)} bytes")
        if doc["plantings"]:
            errors.append(f"{uid}: zero_byte must record no plantings")
    elif kind in ("xlsx_as_pdf", "docx_as_txt"):
        if not data.startswith(b"PK\x03\x04"):
            errors.append(f"{uid}: {kind} does not sniff as an OOXML zip")
        true_name = f"content{'.xlsx' if kind == 'xlsx_as_pdf' else '.docx'}"
        _check_attachment_plantings(true_name, data, doc["plantings"], uid, errors)
    elif kind == "png_as_xlsx":
        if not data.startswith(b"\x89PNG"):
            errors.append(f"{uid}: png_as_xlsx does not sniff as PNG")
        _check_png_image(data, doc["plantings"], uid, errors, stats)
    else:
        errors.append(f"{uid}: unknown problem kind {kind!r}")


# --- per-document dispatch -------------------------------------------------


def _check_document(doc: dict, corpus_dir: Path, errors: list[str],
                    stats: _OcrStats) -> None:
    path = corpus_dir / doc["filename"]
    uid = doc["doc_uid"]
    if not path.exists():
        errors.append(f"{uid}: file missing: {doc['filename']}")
        return

    if doc["problem"] is not None:
        _check_problem(doc, path, errors, stats)
        return

    file_class = doc["file_class"]
    if file_class == "pdf_digital":
        with pymupdf.open(path) as pdf:
            _check_pdf_text(pdf, doc["plantings"], uid, errors)
    elif file_class == "pdf_scanned":
        with pymupdf.open(path) as pdf:
            for page in pdf:
                if page.get_text().strip():
                    errors.append(f"{uid}: scanned pdf has a text layer")
                    break
            _check_pdf_scanned(pdf, doc["plantings"], uid, errors, stats)
    elif file_class == "docx":
        _check_docx(Document(str(path)), doc["plantings"], uid, errors)
    elif file_class == "xlsx":
        wb = load_workbook(path, read_only=True)
        try:
            _check_xlsx(wb, doc["plantings"], uid, errors)
        finally:
            wb.close()
    elif file_class == "eml":
        _check_eml(doc, path, errors)
    elif file_class == "png":
        _check_png_image(path.read_bytes(), doc["plantings"], uid, errors, stats)
    else:  # csv | txt | html — line-oriented plain read
        _check_lines(path.read_text().splitlines(), doc["plantings"], uid, errors)


def _check_scenarios(manifest: dict, errors: list[str]) -> None:
    cfg = PROFILES[manifest["profile"]]
    identities = manifest["identities"]
    documents = manifest["documents"]
    by_uid = {i["person_uid"]: i for i in identities}

    def tagged(tag: str) -> list[dict]:
        return [d for d in documents if tag in d["scenario_tags"]]

    # Every identity appears in at least one planting.
    covered = {
        p["person_uid"] for d in documents for p in d["plantings"] if p["person_uid"]
    }
    for missing in sorted(set(by_uid) - covered):
        errors.append(f"coverage: identity {missing} appears in no document")

    # NicknameCluster: enough persons, 3-5 docs each, ≥3 distinct name forms.
    name_docs: dict[str, list[str]] = defaultdict(list)
    for d in tagged("nickname_cluster"):
        for p in d["plantings"]:
            if p["element_type"] == "name":
                name_docs[p["person_uid"]].append(p["value"])
    if len(name_docs) < cfg.nickname_cluster_persons:
        errors.append(
            f"nickname_cluster: {len(name_docs)} persons < {cfg.nickname_cluster_persons}"
        )
    for uid, names in name_docs.items():
        if not (cfg.nickname_docs_min <= len(names) <= cfg.nickname_docs_max):
            errors.append(f"nickname_cluster: {uid} has {len(names)} docs (want 3-5)")
        if len(set(names)) < 3:
            errors.append(f"nickname_cluster: {uid} uses only {len(set(names))} name forms")

    # SharedName: every canonical-name duplicate is scripted, pairs differ
    # on SSN and DOB, and both members have tagged docs.
    scripted = {frozenset(pair) for pair in manifest["scenario_script"]["shared_name_pairs"]}
    if len(scripted) < cfg.shared_name_pairs:
        errors.append(f"shared_name: {len(scripted)} pairs < {cfg.shared_name_pairs}")
    by_name: dict[str, list[str]] = defaultdict(list)
    for i in identities:
        by_name[i["canonical_name"]].append(i["person_uid"])
    duplicate_groups = {frozenset(uids) for uids in by_name.values() if len(uids) > 1}
    if duplicate_groups != scripted:
        errors.append(
            f"shared_name: duplicate-name groups {duplicate_groups} != scripted {scripted}"
        )
    shared_doc_persons = Counter(
        p["person_uid"]
        for d in tagged("shared_name")
        for p in d["plantings"]
        if p["element_type"] == "name"
    )
    for pair in scripted:
        a, b = tuple(pair)
        if by_uid[a]["elements"]["ssn"] == by_uid[b]["elements"]["ssn"]:
            errors.append(f"shared_name: {a}/{b} share an SSN")
        if by_uid[a]["dob"] == by_uid[b]["dob"]:
            errors.append(f"shared_name: {a}/{b} share a DOB")
        for uid in (a, b):
            if shared_doc_persons[uid] < 2:
                errors.append(f"shared_name: {uid} has <2 tagged docs")

    # PartialIdentifiers: per person, one tagged doc with SSN and no name,
    # one with name (+ last-4) and no full SSN.
    ssn_only: set[str] = set()
    name_only: set[str] = set()
    for d in tagged("partial_identifiers"):
        per_person: dict[str, set[str]] = defaultdict(set)
        for p in d["plantings"]:
            if p["person_uid"]:
                per_person[p["person_uid"]].add(p["element_type"])
        for uid, kinds in per_person.items():
            if "ssn" in kinds and "name" not in kinds:
                ssn_only.add(uid)
            if "name" in kinds and "ssn" not in kinds:
                name_only.add(uid)
    joined = ssn_only & name_only
    if len(joined) < cfg.partial_identifier_persons:
        errors.append(
            f"partial_identifiers: {len(joined)} joinable persons < {cfg.partial_identifier_persons}"
        )

    # BulkSpreadsheet: one xlsx exposing the quota of distinct persons, and
    # its evil-twin PNG re-rendering the same content as an image.
    bulk_docs = tagged("bulk_spreadsheet")

    def bulk_persons(d: dict) -> set[str]:
        return {p["person_uid"] for p in d["plantings"] if p["person_uid"]}

    bulk_xlsx = [d for d in bulk_docs if d["file_class"] == "xlsx"]
    if not any(len(bulk_persons(d)) >= cfg.bulk_spreadsheet_rows for d in bulk_xlsx):
        errors.append(
            f"bulk_spreadsheet: no xlsx doc exposes >= {cfg.bulk_spreadsheet_rows} persons"
        )
    twin = [d for d in bulk_docs if d["file_class"] == "png"]
    if not any(
        len(bulk_persons(d)) >= cfg.bulk_spreadsheet_rows
        and bulk_persons(d) == bulk_persons(x)
        for d in twin
        for x in bulk_xlsx
    ):
        errors.append("bulk_spreadsheet: no evil-twin png mirrors the xlsx sheet's persons")

    # FalsePositiveTraps: doc quota met, every trap kind present, and no
    # trap value collides with any identity element.
    trap_docs = tagged("false_positive_traps")
    if len(trap_docs) < cfg.trap_docs:
        errors.append(f"false_positive_traps: {len(trap_docs)} docs < {cfg.trap_docs}")
    trap_kinds_seen = {
        p["element_type"] for d in trap_docs for p in d["plantings"]
    }
    for kind in TRAP_KINDS:
        if kind not in trap_kinds_seen:
            errors.append(f"false_positive_traps: kind {kind} never planted")
    identity_values = {
        v for i in identities for v in i["elements"].values()
    }
    for d in documents:
        for p in d["plantings"]:
            if p["element_type"].startswith("trap_") and p["value"] in identity_values:
                errors.append(
                    f"{d['doc_uid']}: trap value {p['value']!r} collides with a real identity"
                )

    # ScannedBatch: quota met, plantings presented as image.
    scanned = tagged("scanned_batch")
    if len(scanned) < cfg.scanned_docs:
        errors.append(f"scanned_batch: {len(scanned)} docs < {cfg.scanned_docs}")
    for d in scanned:
        if d["file_class"] != "pdf_scanned":
            errors.append(f"{d['doc_uid']}: scanned_batch doc is {d['file_class']}")
        for p in d["plantings"]:
            if p["presentation"] != "image":
                errors.append(f"{d['doc_uid']}: scanned planting not presentation=image")

    # EmailThreads: quota, attachment class coverage, and the scripted
    # identical-bytes pairs (measured dedup ground truth).
    emails = tagged("email_thread")
    if len(emails) < cfg.eml_docs:
        errors.append(f"email_thread: {len(emails)} docs < {cfg.eml_docs}")
    sha_owners: dict[str, set[str]] = defaultdict(set)
    attachment_exts: set[str] = set()
    for d in emails:
        for a in d["attachments"] or []:
            sha_owners[a["sha256"]].add(d["doc_uid"])
            attachment_exts.add(Path(a["filename"]).suffix)
    for ext in (".docx", ".xlsx", ".pdf"):
        if ext not in attachment_exts:
            errors.append(f"email_thread: no {ext} attachment in the corpus")
    shared_count = sum(1 for owners in sha_owners.values() if len(owners) >= 2)
    if shared_count < cfg.eml_shared_attachment_pairs:
        errors.append(
            f"email_thread: {shared_count} shared-attachment sha256s "
            f"< {cfg.eml_shared_attachment_pairs}"
        )

    # PngScreenshots: quota met.
    screenshots = tagged("png_screenshot")
    if len(screenshots) < cfg.png_docs:
        errors.append(f"png_screenshot: {len(screenshots)} docs < {cfg.png_docs}")

    # ProblemFiles: every kind present per set; problem entries only there.
    problem_docs = [d for d in documents if d["problem"] is not None]
    for d in problem_docs:
        if "problem_files" not in d["scenario_tags"]:
            errors.append(f"{d['doc_uid']}: problem entry outside the problem_files scenario")
    kind_counts = Counter(d["problem"]["kind"] for d in problem_docs)
    expected_kinds = {
        "password_pdf", "truncated_pdf", "zero_byte",
        "xlsx_as_pdf", "docx_as_txt", "png_as_xlsx",
    }
    for kind in expected_kinds:
        if kind_counts[kind] != cfg.problem_sets:
            errors.append(
                f"problem_files: kind {kind} x{kind_counts[kind]} != {cfg.problem_sets}"
            )
    if len(problem_docs) != cfg.problem_docs:
        errors.append(f"problem_files: {len(problem_docs)} docs != {cfg.problem_docs}")

    # Size quotas.
    if len(identities) != cfg.n_identities:
        errors.append(f"identities: {len(identities)} != {cfg.n_identities}")
    if len(documents) < cfg.total_docs:
        errors.append(f"documents: {len(documents)} < {cfg.total_docs}")


def _check_counts(manifest: dict, errors: list[str]) -> None:
    documents = manifest["documents"]
    recomputed = {
        "identities": len(manifest["identities"]),
        "documents": len(documents),
        "plantings": sum(len(d["plantings"]) for d in documents),
        "by_file_class": dict(sorted(Counter(d["file_class"] for d in documents).items())),
        "by_scenario": dict(
            sorted(Counter(t for d in documents for t in d["scenario_tags"]).items())
        ),
    }
    if recomputed != manifest["generated_counts"]:
        errors.append(
            f"generated_counts drift: manifest says {manifest['generated_counts']}, "
            f"recomputed {recomputed}"
        )


def run(manifest_path: Path, corpus_dir: Path) -> int:
    manifest = json.loads(Path(manifest_path).read_text())
    corpus_dir = Path(corpus_dir)
    errors: list[str] = []
    stats = _OcrStats()

    # Corpus dir and manifest must list exactly the same files.
    on_disk = {p.name for p in corpus_dir.iterdir() if p.is_file()}
    listed = {d["filename"] for d in manifest["documents"]}
    for name in sorted(on_disk - listed):
        errors.append(f"unlisted file on disk: {name}")
    for name in sorted(listed - on_disk):
        errors.append(f"manifest file missing on disk: {name}")

    for doc in manifest["documents"]:
        _check_document(doc, corpus_dir, errors, stats)
    _check_scenarios(manifest, errors)
    _check_counts(manifest, errors)

    class_docs = Counter(d["file_class"] for d in manifest["documents"])
    class_plantings: Counter = Counter()
    for d in manifest["documents"]:
        class_plantings[d["file_class"]] += len(d["plantings"])
    print(f"{'file_class':<14}{'docs':>6}{'plantings':>12}")
    for file_class in sorted(class_docs):
        print(f"{file_class:<14}{class_docs[file_class]:>6}{class_plantings[file_class]:>12}")
    print(f"{'TOTAL':<14}{sum(class_docs.values()):>6}{sum(class_plantings.values()):>12}")
    for bucket in sorted(stats.total):
        exact, total = stats.exact[bucket], stats.total[bucket]
        print(
            f"OCR exact recovery [{bucket}]: {exact}/{total} ({exact / total:.1%}) "
            f"— validation bar is fuzz>={OCR_FUZZ_THRESHOLD} per planting"
        )

    if errors:
        print(f"\nVALIDATION FAILED — {len(errors)} error(s):")
        for err in errors:
            print(f"  - {err}")
        return 1
    print(f"\nValidation passed: {sum(class_docs.values())} documents, "
          f"{sum(class_plantings.values())} plantings verified at their recorded locations.")
    return 0
