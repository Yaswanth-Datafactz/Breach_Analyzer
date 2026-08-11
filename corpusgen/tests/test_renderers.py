"""New-renderer invariants (docs/plan.md §8): problem-file properties the
investigator fixtures depend on, the eml attachment round-trip (bytes,
sha256, per-part locations), zip-normalization determinism, the scanned
renderer's image-only contract, and whole-pipeline manifest determinism
across two runs of a tiny profile.
"""

import datetime as dt
import hashlib
import io
import random
from email import policy
from email.parser import BytesParser

import pikepdf
import pymupdf
import pytest
from docx import Document
from faker import Faker
from openpyxl import load_workbook

from corpusgen import templates
from corpusgen.__main__ import generate
from corpusgen.config import MINI, CorpusConfig
from corpusgen.identities import generate_identities
from corpusgen.renderers import (
    DocumentSpec,
    EmailAttachment,
    EmailSpec,
    Plant,
    docx as docx_renderer,
    problem_files,
    scanned_pdf,
)
from corpusgen.renderers.eml import normalize_zip_bytes, render as render_eml

DATE = dt.date(2025, 6, 2)


def _env():
    rng = random.Random(42)
    faker = Faker("en_US")
    faker.seed_instance(rng.getrandbits(64))
    pool = generate_identities(rng, faker, MINI)
    staff = templates.make_staff(rng, faker)
    return rng, pool, staff


def _memo_spec(rng, pool, staff) -> tuple[DocumentSpec, Plant]:
    person = pool.identities[0]
    ssn = Plant(person.person_uid, "ssn", person.elements["ssn"])
    spec = templates.hr_memo(
        rng, staff, DATE, person.canonical_name,
        [Plant(person.person_uid, "name", person.canonical_name), ssn],
    )
    return spec, ssn


def _export_spec(rng, pool) -> DocumentSpec:
    rows = [
        {
            "Name": (p.person_uid, p.canonical_name),
            "SSN": (p.person_uid, p.elements["ssn"]),
            "DOB": (p.person_uid, p.dob),
            "Email": (p.person_uid, p.elements["email"]),
            "Phone": (p.person_uid, p.elements["phone"]),
            "Account": (p.person_uid, p.elements["financial_account"]),
        }
        for p in pool.identities[:3]
    ]
    return templates.customer_export(DATE, rows)


class TestProblemFiles:
    def test_password_pdf_locks_and_unlocks(self, tmp_path):
        rng, pool, staff = _env()
        spec, ssn = _memo_spec(rng, pool, staff)
        path = tmp_path / "locked.pdf"
        plantings, problem = problem_files.password_pdf(spec, path, "mbg-123456")
        assert problem["expected_reason_code"] == "password_protected"
        assert problem["recoverable"] is True and problem["password"] == "mbg-123456"
        with pytest.raises(pikepdf.PasswordError):
            pikepdf.open(path)
        with pymupdf.open(path) as pdf:
            assert pdf.authenticate("mbg-123456")
            page = plantings[-1]["location"]["page"]
            assert ssn.value in pdf[page - 1].get_text()

    def test_truncated_pdf_is_headless_wreck_without_ground_truth(self, tmp_path):
        rng, pool, staff = _env()
        spec, ssn = _memo_spec(rng, pool, staff)
        path = tmp_path / "wreck.pdf"
        plantings, problem = problem_files.truncated_pdf(spec, path)
        data = path.read_bytes()
        assert plantings == []
        assert problem["expected_reason_code"] == "corrupt" and not problem["recoverable"]
        assert data.startswith(b"%PDF") and b"%%EOF" not in data
        assert ssn.value.encode() not in data

    def test_zero_byte(self, tmp_path):
        path = tmp_path / "empty.pdf"
        plantings, problem = problem_files.zero_byte(path)
        assert plantings == [] and problem["expected_reason_code"] == "zero_byte"
        assert path.stat().st_size == 0

    def test_wrong_extension_sniffs_to_true_type(self, tmp_path):
        rng, pool, staff = _env()

        xlsx_path = tmp_path / "sheet.pdf"
        plantings, problem = problem_files.xlsx_as_pdf(_export_spec(rng, pool), xlsx_path)
        assert xlsx_path.read_bytes().startswith(b"PK\x03\x04")
        assert problem["true_class"] == "xlsx" and problem["declared_extension"] == ".pdf"
        wb = load_workbook(io.BytesIO(xlsx_path.read_bytes()))
        loc = plantings[0]["location"]
        assert plantings[0]["value"] in str(wb[loc["sheet"]][loc["cell"]].value)

        memo_spec, _ = _memo_spec(rng, pool, staff)
        docx_path = tmp_path / "memo.txt"
        plantings, problem = problem_files.docx_as_txt(memo_spec, docx_path)
        assert docx_path.read_bytes().startswith(b"PK\x03\x04")
        assert problem["true_class"] == "docx"
        d = Document(io.BytesIO(docx_path.read_bytes()))
        para = plantings[-1]["location"]["paragraph"]
        assert plantings[-1]["value"] in d.paragraphs[para - 1].text

        png_path = tmp_path / "sheet.xlsx"
        plantings, problem = problem_files.png_as_xlsx(_export_spec(rng, pool), png_path)
        assert png_path.read_bytes().startswith(b"\x89PNG")
        assert problem["true_class"] == "png"
        assert all(p["presentation"] == "image" for p in plantings)


class TestEmlRoundTrip:
    def _spec(self, tmp_path) -> tuple[EmailSpec, bytes, Plant]:
        rng, pool, staff = _env()
        person = pool.identities[1]
        claim_spec = templates.medical_claim(
            rng, staff, DATE, person.canonical_name,
            [
                Plant(person.person_uid, "name", person.canonical_name),
                Plant(person.person_uid, "dob", person.dob),
            ],
        )
        inner = tmp_path / "claim.docx"
        inner_plantings = docx_renderer.render(claim_spec, inner)
        content = normalize_zip_bytes(inner.read_bytes())
        attachment = EmailAttachment(
            "claim.docx", content,
            "application", "vnd.openxmlformats-officedocument.wordprocessingml.document",
            [
                {**p, "location": {"part": "attachment:claim.docx", **p["location"]}}
                for p in inner_plantings
            ],
        )
        subject = pool.identities[2]
        phone = Plant(subject.person_uid, "phone", subject.elements["phone"])
        spec = templates.email_message(
            rng, staff, DATE, subject.canonical_name,
            [Plant(subject.person_uid, "name", subject.canonical_name), phone],
            [attachment], signature_contact=True,
        )
        return spec, content, phone

    def test_attachment_and_body_round_trip(self, tmp_path):
        spec, content, phone = self._spec(tmp_path)
        path = tmp_path / "mail.eml"
        plantings = render_eml(spec, path)

        with path.open("rb") as fh:
            msg = BytesParser(policy=policy.default).parse(fh)
        attachments = {a.get_filename(): a.get_payload(decode=True)
                       for a in msg.iter_attachments()}
        assert set(attachments) == {"claim.docx"}
        assert hashlib.sha256(attachments["claim.docx"]).hexdigest() == \
            hashlib.sha256(content).hexdigest()

        body_lines = msg.get_body(preferencelist=("plain",)).get_content().splitlines()
        body_plant = next(p for p in plantings if p["value"] == phone.value)
        assert body_plant["location"]["part"] == "body"
        assert phone.value in body_lines[body_plant["location"]["line"] - 1]

        # Header + signature contact strings all recorded as trap plantings.
        header_plants = [p for p in plantings if p["location"]["part"] == "headers"]
        assert {p["location"]["header"] for p in header_plants} == {"From", "To"}
        for p in header_plants:
            assert p["value"] in str(msg[p["location"]["header"]])
        assert any(p["element_type"] == "trap_signature_email" for p in plantings)

        # Attachment plantings resolve inside the decoded attachment bytes.
        att_plants = [p for p in plantings
                      if p["location"]["part"] == "attachment:claim.docx"]
        d = Document(io.BytesIO(attachments["claim.docx"]))
        for p in att_plants:
            loc = p["location"]
            cell = d.tables[loc["table"] - 1].rows[loc["row"] - 1].cells[loc["col"] - 1]
            assert p["value"] in cell.text

    def test_zip_normalization_is_deterministic(self, tmp_path):
        rng, pool, staff = _env()
        spec, _ = _memo_spec(rng, pool, staff)
        a, b = tmp_path / "a.docx", tmp_path / "b.docx"
        docx_renderer.render(spec, a)
        docx_renderer.render(spec, b)
        assert normalize_zip_bytes(a.read_bytes()) == normalize_zip_bytes(b.read_bytes())


class TestScannedPdf:
    def test_image_only_pdf_with_image_presentation(self, tmp_path):
        rng, pool, staff = _env()
        spec, ssn = _memo_spec(rng, pool, staff)
        path = tmp_path / "scan.pdf"
        plantings = scanned_pdf.render(spec, path)
        assert all(p["presentation"] == "image" for p in plantings)
        assert any(p["value"] == ssn.value for p in plantings)
        with pymupdf.open(path) as pdf:
            for page in pdf:
                assert not page.get_text().strip()  # no text layer
                assert page.get_images()


TINY = CorpusConfig(
    profile="mini",  # reuse mini quotas in validate; only generation runs here
    n_identities=12,
    target_docs=24,
    scanned_docs=2,
    eml_docs=3,
    eml_shared_attachment_pairs=1,
    png_docs=1,
    problem_sets=1,
    nickname_cluster_persons=2,
    nickname_docs_min=3,
    nickname_docs_max=5,
    shared_name_pairs=1,
    partial_identifier_persons=2,
    partial_dump_chunk=4,
    bulk_spreadsheet_rows=4,
    trap_docs=5,
    maiden_fraction=0.3,
)


class TestManifestDeterminism:
    def test_two_runs_same_seed_identical_manifests(self, tmp_path):
        m1, m2 = tmp_path / "m1.json", tmp_path / "m2.json"
        generate(7, TINY, tmp_path / "c1", m1)
        generate(7, TINY, tmp_path / "c2", m2)
        assert m1.read_bytes() == m2.read_bytes()
