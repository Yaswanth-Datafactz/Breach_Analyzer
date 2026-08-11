"""DOCX renderer (python-docx).

Locations: prose plants use {"paragraph": n} — 1-based index into the
document's paragraph sequence (the title heading is paragraph 1); table
plants use {"table": t, "row": r, "col": c} — 1-based, row 1 is the header
row. More precise than a bare line number and directly resolvable by
validate.py through python-docx.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from corpusgen.renderers import FIXED_DOC_DT, DocumentSpec, Paragraph, Table, planting_dict


def render(spec: DocumentSpec, path: Path) -> list[dict]:
    doc = Document()
    plantings: list[dict] = []
    para_count = 0
    table_count = 0

    doc.add_heading(spec.title, level=1)
    para_count += 1
    doc.add_paragraph(spec.doc_date.isoformat())
    para_count += 1

    for block in spec.blocks:
        if isinstance(block, Paragraph):
            lines = block.text.split("\n")
            line_para_idx: list[int] = []
            for line in lines:
                doc.add_paragraph(line)
                para_count += 1
                line_para_idx.append(para_count)
            for plant in block.plants:
                idx = next(
                    (line_para_idx[i] for i, line in enumerate(lines) if plant.value in line),
                    None,
                )
                assert idx is not None, f"plant {plant.value!r} not in paragraph text"
                plantings.append(planting_dict(plant, {"paragraph": idx}))
        else:
            assert isinstance(block, Table)
            table = doc.add_table(rows=1 + len(block.rows), cols=len(block.headers))
            table.style = "Table Grid"
            table_count += 1
            for c, header in enumerate(block.headers):
                table.rows[0].cells[c].text = header
            for r, row in enumerate(block.rows):
                for c, value in enumerate(row):
                    table.rows[r + 1].cells[c].text = value
            for row_idx, col_idx, plant in block.cell_plants:
                assert plant.value in block.rows[row_idx][col_idx]
                plantings.append(
                    planting_dict(
                        plant,
                        {"table": table_count, "row": row_idx + 2, "col": col_idx + 1},
                    )
                )

    props = doc.core_properties
    props.created = FIXED_DOC_DT
    props.modified = FIXED_DOC_DT
    doc.save(str(path))
    return plantings
