"""DOCX parsing via python-docx.

Locator vocabulary mirrors corpusgen/renderers/docx.py exactly: prose is
{"paragraph": n}, 1-based over the document's top-level paragraph sequence
(the title heading is paragraph 1 -- python-docx's `doc.paragraphs` yields
that same sequence, excluding table cells); tables are {"table": t, ...},
1-based over `doc.tables`. A manifest planting and the passage holding it
therefore share coordinates without translation.

Granularity: one passage per non-empty paragraph, one passage per TABLE
(not per cell): a Field/Value claim table splits the member's name and SSN
across rows, and per-row passages would sever exactly the name<->identifier
adjacency mention-building (tier 1) and evidence anchoring need. The table
passage's text renders row 1 (headers) first, then each data row, cells
joined with " | " -- cell values stay byte-exact so `passage.find(value)`
resolves plantings and computed offsets (plan D9) alike.
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument

from app.services.parsing.types import ParsedPassage, ParseResult

_CELL_SEPARATOR = " | "


def parse_docx(content: bytes) -> ParseResult:
    """Raises on corrupt input (a truncated docx fails in python-docx's zip
    layer) -- the pipeline's per-document exception boundary maps that to a
    `corrupt` quarantine; this module never swallows it."""
    doc = DocxDocument(io.BytesIO(content))
    passages: list[ParsedPassage] = []

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        if not paragraph.text.strip():
            continue  # numbering stays true to document order; empties just emit nothing
        passages.append(
            ParsedPassage(kind="text_block", locator={"paragraph": idx}, text=paragraph.text)
        )

    for table_idx, table in enumerate(doc.tables, start=1):
        rows = table.rows
        if not rows:
            continue
        lines = [
            _CELL_SEPARATOR.join(cell.text for cell in row.cells) for row in rows
        ]
        passages.append(
            ParsedPassage(
                kind="text_block",
                # row 1 is the header row (corpusgen's convention); data
                # rows are 2..len(rows), matching manifest {"table","row","col"}.
                locator={"table": table_idx, "row_start": 1, "row_end": len(rows)},
                text="\n".join(lines),
            )
        )

    return ParseResult(passages=passages)
